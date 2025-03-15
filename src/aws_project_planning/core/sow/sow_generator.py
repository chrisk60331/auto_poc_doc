"""Statement of Work (SOW) generator module."""

import os
import shutil
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt
from jinja2 import Environment, FileSystemLoader


@dataclass
class SOWSection:
    """Represents a section in the Statement of Work."""

    title: str
    content: str
    subsections: Optional[List["SOWSection"]] = None


@dataclass
class SOWData:
    """Data structure for SOW content."""

    project_name: str
    client_name: str
    start_date: datetime
    end_date: datetime
    project_description: str
    scope: List[str]
    deliverables: List[str]
    assumptions: List[str]
    timeline: List[Dict[str, Any]]
    cost: Dict[str, Any]


class SOWGenerator:
    """Generates Statement of Work documents."""

    def __init__(self, template_dir: str = "templates/sow"):
        """Initialize SOW generator with template directory."""
        self.template_dir = Path(template_dir)
        self.env = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    def _create_document(self) -> Document:
        """Create a new Word document with basic styling."""
        doc = Document()
        
        # Set up default styles
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        
        # Set up margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        return doc

    def _add_header(self, doc: Document, data: SOWData) -> None:
        """Add header section to the document."""
        doc.add_heading("Statement of Work", 0)
        doc.add_paragraph(f"Project: {data.project_name}")
        doc.add_paragraph(f"Client: {data.client_name}")
        doc.add_paragraph(f"Date: {data.start_date.strftime('%B %d, %Y')}")

    def _add_section(self, doc: Document, section: SOWSection, level: int = 1) -> None:
        """Add a section to the document."""
        doc.add_heading(section.title, level)
        doc.add_paragraph(section.content)
        
        if section.subsections:
            for subsection in section.subsections:
                self._add_section(doc, subsection, level + 1)

    def generate(self, data: SOWData, template_name: str, output_path: str) -> None:
        """Generate SOW document using provided data and template."""
        # Check if docx template exists
        docx_template_path = self.template_dir / f"{template_name}.docx"
        if docx_template_path.exists():
            # Use docx template directly
            self._generate_from_docx_template(data, docx_template_path, output_path)
            return
            
        # Fallback to Jinja2 template if docx template doesn't exist
        j2_template_path = self.template_dir / f"{template_name}.j2"
        if j2_template_path.exists():
            # Generate from Jinja2 template
            self._generate_from_j2_template(data, template_name, output_path)
            return
            
        # If neither template exists, raise an error
        raise ValueError(f"Template file not found for {template_name}")

    def _generate_from_j2_template(self, data: SOWData, template_name: str, output_path: str) -> None:
        """Generate SOW document using a Jinja2 template."""
        # Load template
        template = self.env.get_template(f"{template_name}.j2")
        
        # Create document
        doc = self._create_document()
        
        # Add header
        self._add_header(doc, data)
        
        # Add project overview
        overview = SOWSection(
            title="Project Overview",
            content=data.project_description
        )
        self._add_section(doc, overview)
        
        # Add scope
        scope_content = "\n".join([f"- {item}" for item in data.scope])
        scope = SOWSection(
            title="Scope of Work",
            content=scope_content
        )
        self._add_section(doc, scope)
        
        # Add deliverables
        deliverables_content = "\n".join([f"- {item}" for item in data.deliverables])
        deliverables = SOWSection(
            title="Deliverables",
            content=deliverables_content
        )
        self._add_section(doc, deliverables)
        
        # Add timeline
        timeline_content = "\n".join(
            [f"- {item['phase']}: {item['duration']}" for item in data.timeline]
        )
        timeline = SOWSection(
            title="Project Timeline",
            content=timeline_content
        )
        self._add_section(doc, timeline)
        
        # Add cost
        cost_content = (
            f"Total Project Cost: ${data.cost['total']:,.2f}\n"
            f"Payment Schedule:\n"
            + "\n".join([f"- {k}: ${v:,.2f}" for k, v in data.cost['schedule'].items()])
        )
        cost = SOWSection(
            title="Cost and Payment Schedule",
            content=cost_content
        )
        self._add_section(doc, cost)
        
        # Add assumptions
        assumptions_content = "\n".join([f"- {item}" for item in data.assumptions])
        assumptions = SOWSection(
            title="Assumptions and Constraints",
            content=assumptions_content
        )
        self._add_section(doc, assumptions)
        
        # Save document
        doc.save(output_path)
        
    def _generate_from_docx_template(self, data: SOWData, template_path: str, output_path: str) -> None:
        """Generate SOW document using a .docx template file."""
        # Make a copy of the template
        shutil.copy(template_path, output_path)
        
        # Open the copied template for modification
        doc = Document(output_path)
        
        # Set prepared date to today and effective date to 2 weeks from today
        prepared_date = datetime.now()
        effective_date = prepared_date + timedelta(days=14)
        prepared_date_str = prepared_date.strftime('%B %d, %Y')
        effective_date_str = effective_date.strftime('%B %d, %Y')
        
        # We'll need to find and replace placeholders in the document
        # Common placeholders to replace with additional variations for client/customer name
        replacements = {
            "{{PROJECT_NAME}}": data.project_name,
            "{{CLIENT_NAME}}": data.client_name,
            "{{CUSTOMER_NAME}}": data.client_name,  # Add CUSTOMER_NAME variant
            "{{CUSTOMER}}": data.client_name,       # Add CUSTOMER variant
            "{{CLIENT}}": data.client_name,         # Add CLIENT variant
            "{{PROJECT_DESCRIPTION}}": data.project_description,
            "{{OVERVIEW}}": data.project_description,  # Add OVERVIEW variant
            "{{START_DATE}}": data.start_date.strftime('%B %d, %Y'),
            "{{END_DATE}}": data.end_date.strftime('%B %d, %Y'),
            "{{PREPARED_DATE}}": prepared_date_str,
            "{{EFFECTIVE_DATE}}": effective_date_str,
            "{{TOTAL_COST}}": f"${data.cost['total']:,.2f}",
        }
        
        # Special handling for lists and complex data
        scope_content = "\n".join([f"• {item}" for item in data.scope])
        deliverables_content = "\n".join([f"• {item}" for item in data.deliverables])
        assumptions_content = "\n".join([f"• {item}" for item in data.assumptions])
        timeline_content = "\n".join([f"• {item['phase']}: {item['duration']}" for item in data.timeline])
        cost_schedule = "\n".join([f"• {k}: ${v:,.2f}" for k, v in data.cost['schedule'].items()])
        
        replacements.update({
            "{{SCOPE}}": scope_content,
            "{{DELIVERABLES}}": deliverables_content,
            "{{ASSUMPTIONS}}": assumptions_content,
            "{{TIMELINE}}": timeline_content,
            "{{PAYMENT_SCHEDULE}}": cost_schedule,
        })
        
        # Add more variations of placeholders to handle different formats
        additional_placeholders = {}
        for key, value in replacements.items():
            # Try alternative formats
            clean_key = key.replace("{{", "").replace("}}", "")
            additional_placeholders[f"{{{clean_key}}}"] = value      # {PLACEHOLDER}
            additional_placeholders[clean_key] = value                # PLACEHOLDER
            additional_placeholders[f"<{clean_key}>"] = value         # <PLACEHOLDER>
            additional_placeholders[f"[{clean_key}]"] = value         # [PLACEHOLDER]
            additional_placeholders[f"{clean_key.lower()}"] = value   # lowercase
            additional_placeholders[f"{clean_key.upper()}"] = value   # UPPERCASE
            
        replacements.update(additional_placeholders)
        
        # Additional specific date placeholders
        date_placeholders = {
            "PREPARED DATE": prepared_date_str,
            "EFFECTIVE DATE": effective_date_str,
            "PREPARED": prepared_date_str,
            "EFFECTIVE": effective_date_str,
        }
        replacements.update(date_placeholders)
        
        # Special specific handling for Overview/Project Description and Customer Name
        # Try to find any paragraphs that look like they contain these sections
        overview_keywords = ["overview", "project overview", "description", "project description", "summary", "project summary"]
        customer_keywords = ["customer", "client", "customer name", "client name"]
        date_keywords = {
            "prepared date": prepared_date_str,
            "effective date": effective_date_str,
            "prepared": prepared_date_str,
            "effective": effective_date_str,
        }
        
        for paragraph in doc.paragraphs:
            # Check for overview section indicators
            lower_text = paragraph.text.lower()
            
            # Look for date keywords
            for date_key, date_value in date_keywords.items():
                if date_key in lower_text:
                    # Try different replacement strategies
                    if ":" in paragraph.text:
                        # Format: "Prepared Date: [date]"
                        parts = paragraph.text.split(":", 1)
                        if len(parts) > 1:
                            new_text = f"{parts[0]}: {date_value}"
                            paragraph.clear()
                            paragraph.add_run(new_text)
                    else:
                        # Try to replace the text after the keyword
                        key_index = lower_text.find(date_key)
                        if key_index != -1:
                            prefix = paragraph.text[:key_index + len(date_key)]
                            # Check if there's text after the keyword
                            if len(paragraph.text) > key_index + len(date_key):
                                suffix = paragraph.text[key_index + len(date_key):]
                                # Check if suffix contains a date pattern
                                import re
                                date_pattern = re.compile(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\w+ \d{1,2},? \d{4}')
                                match = date_pattern.search(suffix)
                                if match:
                                    # Replace the date pattern with our date
                                    new_suffix = date_pattern.sub(date_value, suffix)
                                    new_text = prefix + new_suffix
                                else:
                                    # Just add our date
                                    new_text = prefix + " " + date_value
                            else:
                                # No text after keyword, just add our date
                                new_text = prefix + " " + date_value
                            paragraph.clear()
                            paragraph.add_run(new_text)
            
            # Look for overview section indicators
            if any(keyword in lower_text for keyword in overview_keywords):
                # Either replace in this paragraph or the next
                if ":" in paragraph.text or paragraph.text.strip().endswith(":"):
                    # This paragraph ends with a colon, so the next paragraph is likely the content
                    paragraph_index = [p.text for p in doc.paragraphs].index(paragraph.text)
                    if paragraph_index < len(doc.paragraphs) - 1:
                        next_paragraph = doc.paragraphs[paragraph_index + 1]
                        next_paragraph.clear()
                        next_paragraph.add_run(data.project_description)
                else:
                    # Try to replace content in this paragraph
                    new_text = paragraph.text
                    for key_word in overview_keywords:
                        if key_word in lower_text:
                            # Try to replace after this keyword or its variations
                            pattern = f"{key_word}:"
                            if pattern in lower_text:
                                parts = paragraph.text.split(":", 1)
                                if len(parts) > 1:
                                    new_text = f"{parts[0]}: {data.project_description}"
                                    paragraph.clear()
                                    paragraph.add_run(new_text)
                                    break
            
            # Look for customer/client name indicators
            if any(keyword in lower_text for keyword in customer_keywords):
                # Try different replacement strategies
                if ":" in paragraph.text:
                    # Format: "Customer: [name]"
                    parts = paragraph.text.split(":", 1)
                    if len(parts) > 1:
                        new_text = f"{parts[0]}: {data.client_name}"
                        paragraph.clear()
                        paragraph.add_run(new_text)
                elif "name" in lower_text:
                    # Look for "customer name" or similar patterns
                    for keyword in customer_keywords:
                        if keyword in lower_text:
                            # Replace after the keyword
                            keyword_index = lower_text.find(keyword)
                            if keyword_index != -1:
                                prefix = paragraph.text[:keyword_index + len(keyword)]
                                new_text = f"{prefix} {data.client_name}"
                                paragraph.clear()
                                paragraph.add_run(new_text)
                                break
        
        # Enhanced text replacement for paragraphs
        for paragraph in doc.paragraphs:
            replaced = False
            original_text = paragraph.text
            
            # Check for complete placeholder paragraph matches
            for key, value in replacements.items():
                if original_text.strip() == key.strip():
                    # Instead of just replacing text, we need to clear and rebuild the paragraph
                    # because simply setting text might lose formatting
                    paragraph.clear()
                    run = paragraph.add_run(value)
                    replaced = True
                    break
            
            # If paragraph wasn't a complete match, check for placeholder in text
            if not replaced:
                new_text = original_text
                for key, value in replacements.items():
                    if key in new_text:
                        new_text = new_text.replace(key, value)
                
                if new_text != original_text:
                    # Replace text while trying to preserve formatting
                    paragraph.clear()
                    paragraph.add_run(new_text)
        
        # Process tables with enhanced handling
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        new_text = original_text
                        
                        for key, value in replacements.items():
                            if key in new_text:
                                new_text = new_text.replace(key, value)
                        
                        if new_text != original_text:
                            # Replace text while trying to preserve formatting
                            paragraph.clear()
                            paragraph.add_run(new_text)
        
        # Additional handling for document properties and headers/footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                original_text = header.text
                new_text = original_text
                
                for key, value in replacements.items():
                    if key in new_text:
                        new_text = new_text.replace(key, value)
                
                if new_text != original_text:
                    header.clear()
                    header.add_run(new_text)
            
            for footer in section.footer.paragraphs:
                original_text = footer.text
                new_text = original_text
                
                for key, value in replacements.items():
                    if key in new_text:
                        new_text = new_text.replace(key, value)
                
                if new_text != original_text:
                    footer.clear()
                    footer.add_run(new_text)
        
        # Save the document
        doc.save(output_path) 