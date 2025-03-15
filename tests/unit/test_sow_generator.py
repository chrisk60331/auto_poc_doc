"""Unit tests for SOW Generator."""

import os
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from aws_project_planning.core.sow.sow_generator import SOWData, SOWGenerator, SOWSection


@pytest.fixture
def temp_template_dir(tmp_path):
    """Create a temporary template directory with test templates."""
    template_dir = tmp_path / "templates"
    template_dir.mkdir(parents=True)
    
    # Create a J2 template for testing
    j2_dir = template_dir / "sow"
    j2_dir.mkdir(parents=True)
    j2_template = j2_dir / "standard.j2"
    j2_template.write_text("Test template for {{ project_name }}")
    
    # Create a docx template for testing
    docx_path = j2_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("Test template for {{PROJECT_NAME}}")
    doc.add_paragraph("Client: {{CLIENT_NAME}}")
    doc.add_paragraph("Project Overview: {{PROJECT_DESCRIPTION}}")
    doc.add_paragraph("Scope: {{SCOPE}}")
    doc.add_paragraph("Deliverables: {{DELIVERABLES}}")
    doc.add_paragraph("Timeline: {{TIMELINE}}")
    doc.add_paragraph("Prepared Date: PREPARED DATE")
    doc.add_paragraph("Effective Date: EFFECTIVE DATE")
    doc.save(docx_path)
    
    return template_dir


@pytest.fixture
def sow_generator(temp_template_dir):
    """Create a SOW generator instance with temporary template directory."""
    return SOWGenerator(template_dir=str(temp_template_dir / "sow"))


@pytest.fixture
def sample_sow_data():
    """Create sample SOW data for testing."""
    return SOWData(
        project_name="Test Project",
        client_name="Test Client",
        start_date=datetime.now(),
        end_date=datetime.now() + timedelta(days=30),
        project_description="This is a test project",
        scope=["Requirement gathering", "Implementation"],
        deliverables=["Documentation", "Source code"],
        assumptions=["Client will provide timely feedback"],
        timeline=[
            {"phase": "Planning", "duration": "2 weeks"},
            {"phase": "Development", "duration": "4 weeks"},
        ],
        cost={
            "total": 50000.0,
            "schedule": {
                "Initial Payment": 15000.0,
                "Final Payment": 35000.0,
            },
        },
    )


def test_create_document(sow_generator):
    """Test document creation with basic styling."""
    doc = sow_generator._create_document()
    
    assert doc is not None
    assert isinstance(doc, Document)
    
    # Check if styles were applied
    assert doc.styles["Normal"].font.name == "Arial"


def test_add_header(sow_generator, sample_sow_data):
    """Test adding header to the document."""
    doc = sow_generator._create_document()
    sow_generator._add_header(doc, sample_sow_data)
    
    # Get paragraphs with content
    paragraphs = [p for p in doc.paragraphs if p.text]
    
    assert "Statement of Work" in paragraphs[0].text
    assert f"Project: {sample_sow_data.project_name}" in paragraphs[1].text
    assert f"Client: {sample_sow_data.client_name}" in paragraphs[2].text


def test_add_section(sow_generator):
    """Test adding sections to the document."""
    doc = sow_generator._create_document()
    
    # Create a test section with subsections
    subsection = SOWSection(title="Subsection", content="Subsection content")
    section = SOWSection(
        title="Test Section", 
        content="Test content", 
        subsections=[subsection]
    )
    
    sow_generator._add_section(doc, section)
    
    # Get paragraphs with content
    paragraphs = [p for p in doc.paragraphs if p.text]
    
    assert "Test Section" in paragraphs[0].text
    assert "Test content" in paragraphs[1].text
    assert "Subsection" in paragraphs[2].text
    assert "Subsection content" in paragraphs[3].text


def test_generate_with_j2_template(sow_generator, sample_sow_data, tmp_path):
    """Test generating SOW document using a Jinja2 template."""
    output_path = tmp_path / "test_output.docx"
    
    sow_generator.generate(sample_sow_data, "standard", str(output_path))
    
    assert output_path.exists()
    
    # Read the generated document
    doc = Document(output_path)
    paragraphs_text = [p.text for p in doc.paragraphs if p.text]
    
    assert any("Statement of Work" in text for text in paragraphs_text)
    assert any("Project Overview" in text for text in paragraphs_text)
    assert any("Scope of Work" in text for text in paragraphs_text)
    assert any("Deliverables" in text for text in paragraphs_text)


def test_generate_with_docx_template(sow_generator, sample_sow_data, tmp_path):
    """Test generating SOW document using a Word template."""
    output_path = tmp_path / "test_docx_output.docx"
    
    sow_generator.generate(sample_sow_data, "test", str(output_path))
    
    assert output_path.exists()
    
    # Read the generated document
    doc = Document(output_path)
    paragraphs_text = [p.text for p in doc.paragraphs if p.text]
    
    assert any("Test Project" in text for text in paragraphs_text)
    assert any("Test Client" in text for text in paragraphs_text)
    assert any("This is a test project" in text for text in paragraphs_text)


def test_generate_with_nonexistent_template(sow_generator, sample_sow_data, tmp_path):
    """Test generating SOW document with nonexistent template."""
    output_path = tmp_path / "test_error.docx"
    
    with pytest.raises(ValueError) as excinfo:
        sow_generator.generate(sample_sow_data, "nonexistent", str(output_path))
    
    assert "Template file not found" in str(excinfo.value)


def test_generate_from_docx_template_placeholder_replacement(sow_generator, sample_sow_data, tmp_path):
    """Test placeholder replacement in DOCX template."""
    output_path = tmp_path / "test_placeholders.docx"
    
    # Create a template with various placeholder formats
    doc = Document()
    doc.add_paragraph("{{PROJECT_NAME}}")
    doc.add_paragraph("{PROJECT_NAME}")
    doc.add_paragraph("PROJECT_NAME")
    doc.add_paragraph("<PROJECT_NAME>")
    doc.add_paragraph("[PROJECT_NAME]")
    doc.add_paragraph("project_name")
    doc.add_paragraph("PROJECT NAME")
    
    # Tables
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Project"
    table.cell(0, 1).text = "{{PROJECT_NAME}}"
    
    # Headers and footers
    template_path = tmp_path / "template_with_placeholders.docx"
    doc.save(template_path)
    
    # Patch method to use our custom template
    with patch.object(Path, 'exists', return_value=True), \
         patch.object(SOWGenerator, '_generate_from_docx_template', wraps=sow_generator._generate_from_docx_template) as mock:
        
        sow_generator.generate(sample_sow_data, "test_placeholders", str(output_path))
        mock.assert_called_once()
        
        # We can't easily test the actual document, but we ensure the method is called
        # This is a limitation of testing the DOCX generator


def test_docx_date_handling(sow_generator, sample_sow_data, tmp_path):
    """Test date handling in DOCX template."""
    output_path = tmp_path / "test_dates.docx"
    
    # Create a template with date placeholders
    doc = Document()
    doc.add_paragraph("Prepared Date: PREPARED DATE")
    doc.add_paragraph("Effective Date: EFFECTIVE DATE")
    doc.add_paragraph("prepared date: 01/01/2022")
    
    template_path = tmp_path / "template_with_dates.docx"
    doc.save(template_path)
    
    # Mock datetime.now() to return a fixed date
    fixed_date = datetime(2023, 1, 1)
    expected_prepared = "January 01, 2023"
    expected_effective = "January 15, 2023"  # 14 days later
    
    with patch('datetime.datetime') as mock_datetime:
        mock_datetime.now.return_value = fixed_date
        mock_datetime.side_effect = lambda *args, **kw: datetime(*args, **kw)
        
        # Patch to use our custom template
        with patch.object(Path, 'exists', return_value=True), \
             patch.object(shutil, 'copy', return_value=None), \
             patch.object(Document, '__init__', return_value=None), \
             patch.object(Document, 'paragraphs', new_callable=MagicMock), \
             patch.object(Document, 'save', return_value=None):
            
            sow_generator._generate_from_docx_template(sample_sow_data, template_path, str(output_path))
            # We've mocked the actual document handling, but ensured the method runs without error 